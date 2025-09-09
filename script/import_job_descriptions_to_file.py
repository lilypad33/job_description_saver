import os
import re
import pyperclip
from datetime import datetime
from dotenv import load_dotenv

# === SETTINGS ===
DEBUG = True
DEFAULT_FORMAT = "txt"  # "txt" or "docx" (you can still pick at runtime)

# === LOAD ENVIRONMENT VARIABLES ===
load_dotenv()
SAVE_FOLDER = os.getenv("SAVE_FOLDER")
if not SAVE_FOLDER:
    print("Error: SAVE_FOLDER not set in .env")
    raise SystemExit(1)
os.makedirs(SAVE_FOLDER, exist_ok=True)

# === GET CLIPBOARD CONTENT ===
job_description = pyperclip.paste()
if not job_description.strip():
    print("Clipboard is empty. Copy a job description first.")
    raise SystemExit(1)

# === CONSTANTS ===
ROLE_KEYWORDS = [
    "Engineer","Developer","Manager","Analyst","Designer","Scientist","Architect",
    "Administrator","Specialist","Consultant","Coordinator","Lead","Director","Intern",
    "Operator","Technician","Writer","Producer","Editor","Marketer","Researcher",
    "Software","Data","Backend","Frontend","Full Stack","Machine Learning","ML","AI",
    "Security","Cloud","DevOps","SRE","Product","Project","Program","QA","Quality",
    "Support","Sales","Marketing","Finance","HR","People","Operations","IT","UX","UI"
]
PREFERRED_TITLES = {
    "software developer","software engineer","data analyst","data scientist",
    "backend developer","frontend developer","full stack developer","full stack engineer",
    "web developer","application developer", "cloud developer"
}
BANNED_TITLES = {
    "seniority","seniority level","employment","employment type",
    "job","title","position","role","location","locations",
    "department","division","schedule","shift","vacancy"
}

EMAIL_OR_URL = re.compile(r"(https?://\S+|\b\w+@\w+\.\w+)", re.IGNORECASE)
COMMON_FILLER_START = re.compile(r"^(the|a|an|our|we|i)\b", re.IGNORECASE)
TRAILING_DEPT = re.compile(r"\b(team|department|group|program|studio|lab|labs)\b\.?$", re.IGNORECASE)
LEADING_COMPANY_LABEL = re.compile(r"^(company|employer|organization|org)\s*[:\-–]\s*", re.IGNORECASE)
NON_NAME_VERBS = re.compile(r"\b(is|are|seeking|hiring|looking|need|needs|join|build|help|drive|lead)\b", re.IGNORECASE)

def debug(msg):
    if DEBUG:
        print(f"[DEBUG] {msg}")

def sanitize_filename(name):
    return "".join(c for c in name if c.isalnum() or c in (" ", "-", "_")).rstrip()

def looks_title_cased(text):
    small = {"of","and","for","to","in","on","with","the","a","an","or"}
    words = [w for w in re.split(r"\s+", text.strip()) if w]
    if not words:
        return False
    caps = 0
    for i, w in enumerate(words):
        if re.match(r"^[A-Za-z]+$", w):
            if w[0].isupper() or (i > 0 and w.lower() in small):
                caps += 1
        else:
            caps += 1
    return caps >= max(1, int(0.6 * len(words)))

def is_probable_location(text):
    """
    Returns True if the given text looks like a location (city, state, country).
    Designed to filter out location strings from job title or company name fields.
    """

    text = text.strip()
    text_lower = text.lower()

    # US state abbreviations
    state_abbr = {
        "al","ak","az","ar","ca","co","ct","de","fl","ga","hi","id","il","in","ia","ks","ky","la",
        "me","md","ma","mi","mn","ms","mo","mt","ne","nv","nh","nj","nm","ny","nc","nd","oh","ok",
        "or","pa","ri","sc","sd","tn","tx","ut","vt","va","wa","wv","wi","wy"
    }

    # Common country abbreviations (ISO alpha-2)
    country_abbr = {
        "us","uk","ca","au","nz","de","fr","es","it","nl","se","no","fi","ch","jp","cn","in","br","mx"
    }

    # Full US state names
    state_names = {
        "alabama","alaska","arizona","arkansas","california","colorado","connecticut","delaware",
        "florida","georgia","hawaii","idaho","illinois","indiana","iowa","kansas","kentucky",
        "louisiana","maine","maryland","massachusetts","michigan","minnesota","mississippi",
        "missouri","montana","nebraska","nevada","new hampshire","new jersey","new mexico",
        "new york","north carolina","north dakota","ohio","oklahoma","oregon","pennsylvania",
        "rhode island","south carolina","south dakota","tennessee","texas","utah","vermont",
        "virginia","washington","west virginia","wisconsin","wyoming"
    }

    # Common country names
    country_names = {
        "united states","usa","canada","australia","united kingdom","england","scotland","wales",
        "ireland","germany","france","spain","italy","netherlands","sweden","norway","denmark",
        "finland","switzerland","japan","china","india","brazil","mexico"
    }

    # --- Strong comma rule ---
    if "," in text:
        parts = [p.strip().lower() for p in text.split(",", 1)]
        if len(parts) == 2:
            second = parts[1]
            if second in state_abbr or second in country_abbr or second in state_names or second in country_names:
                return True

    # --- All-words rule ---
    words = [w.strip(",.") for w in text_lower.split()]
    if all(w in state_abbr or w in country_abbr or w in state_names or w in country_names for w in words):
        return True

    return False

def find_company_in_next_lines(lines, start_index, max_ahead=5):
    for i in range(start_index + 1, min(start_index + 1 + max_ahead, len(lines))):
        ln = lines[i].strip()
        if not ln:
            continue
        # Match "Company is ..." or "About Company"
        m = re.match(r"^([A-Z][\w& ]+?)\s+is\s+.+$", ln)
        if m and not is_probable_location(m.group(1)):
            return m.group(1).strip()
    return None

def extract_company_and_title_from_line(line):
    # 1. Company is hiring a Job Title
    m = re.match(r"^([A-Z][\w& ]+?)\s+is\s+hiring\s+(?:an?\s+|for\s+an?\s+)?(.+)$", line, flags=re.IGNORECASE)
    if m:
        return m.group(1).strip(), m.group(2).strip(), False

    # 2. Company is seeking a Job Title
    m = re.match(r"^([A-Z][\w& ]+?)\s+is\s+seeking\s+(?:an?\s+)?(.+)$", line, flags=re.IGNORECASE)
    if m:
        return m.group(1).strip(), m.group(2).strip(), False

    # 3. Company is a ... (no title)
    m = re.match(r"^([A-Z][\w& ]+?)\s+is\s+a\s+.+$", line, flags=re.IGNORECASE)
    if m:
        return m.group(1).strip(), None, False

    # 4. Join Company as Job Title
    m = re.match(r"^Join\s+([A-Z][\w& ]+?)\s+as\s+(.+)$", line, flags=re.IGNORECASE)
    if m:
        company = m.group(1).strip()
        title = m.group(2).strip()
        if company.lower() in {"us", "our", "our team"}:
            return None, title, True  # Lookahead needed
        return company, title, False

    return None, None, False

def clean_title(title):
    # --- Remove known stop phrases ---
    stop_phrases = [
        r"\s+who\s+is\s+.*", r"\s+that\s+is\s+.*", r"\s+with\s+experience\s+.*",
        r"\s+to\s+join\s+our\s+team.*", r"\s+to\s+help\s+.*", r"\s+as\s+part\s+of\s+.*",
        r"\s+needed\s+.*", r"\s+ASAP.*", r"\s+immediately.*", r"\s+based\s+in\s+.*",
        r"\s+in\s+[^/\\,:;–—-]+$",
        r"\s+to\s+work\s+onsite.*",
        r"\s+to\s+support\s+.*",
        r"\s+to\s+build\s+.*",
        r"\s+to\s+develop\s+.*",
        r"\s+to\s+design\s+.*"
    ]
    for pat in stop_phrases:
        title = re.sub(pat, "", title, flags=re.IGNORECASE)

    # --- Remove trailing location/company fragments ---
    title = re.sub(r"\s+(at|with|for|in)\s+.+$", "", title, flags=re.IGNORECASE)
    title = re.sub(r",\s*[A-Za-z\s]+,\s*[A-Z]{2}$", "", title)       # , City, ST
    title = re.sub(r",\s*[A-Za-z\s]+,\s*[A-Za-z]+$", "", title)      # , City, Country

    # --- Remove trailing filler like "to join our team" ---
    title = re.sub(r"\s+to\s+join.*$", "", title, flags=re.IGNORECASE)

    # --- Filter out generic non-title phrases ---
    generic_phrases = {
        "member of the team", "part of the team", "part of our team",
        "member of our company", "part of the company", "team member"
    }
    if title.strip().lower() in generic_phrases:
        return ""

    # --- Cut if punctuation followed by lowercase descriptive text ---
    title = re.sub(r"\s*[-–—:]\s+[a-z].*", "", title)

    # --- Remove leading labels ---
    title = re.sub(r"^(Hiring|Role|Position|Title)\s*[:\-–]\s*", "", title, flags=re.IGNORECASE)

    # --- Cut at sentence punctuation if long ---
    title = re.split(r"[.|;]\s", title)[0]

    # --- Limit words ---
    words = title.split()
    if len(words) > 8:
        title = " ".join(words[:8])

    # --- Final cleanup ---
    title = title.strip(" \t-–—:,.").replace("  ", " ")

    # --- Extra: strip trailing location if still present ---
    parts = title.split()
    for cut in range(len(parts), 0, -1):
        if is_probable_location(" ".join(parts[cut-1:])):
            title = " ".join(parts[:cut-1])
            break

    return title


def clean_company(name):
    name = EMAIL_OR_URL.sub("", name)
    name = name.replace("®", "").replace("™", "").replace("©", "")
    name = LEADING_COMPANY_LABEL.sub("", name).strip()
    # LinkedIn artifacts
    name = re.sub(r"\blogo\b", "", name, flags=re.IGNORECASE)
    name = re.sub(r"\bsee jobs?\b", "", name, flags=re.IGNORECASE)
    name = re.sub(r"\bview profile\b", "", name, flags=re.IGNORECASE)
    # Cut at separators if description follows
    name = re.sub(r"\s*[-–—:|]\s+[a-z].*", "", name)
    # Remove trailing department-like words
    name = TRAILING_DEPT.sub("", name).strip()
    # Normalize
    name = re.sub(r"\s+", " ", name).strip(" \t-–—:,.")
    # Keep first chunk before dividers
    name = re.split(r"[|,/]", name)[0].strip()
    return name

def has_role_keyword(text):
    t = text.lower()
    return any(kw.lower() in t for kw in ROLE_KEYWORDS)

# === TITLE CANDIDATES ===
title_candidates = []

def add_title_candidate(val, score, why):
    val = re.sub(r"\s+", " ", val).strip()
    if not val:
        return
    title_candidates.append((val, score, why))
    debug(f"Title candidate (+{score}): '{val}' via {why}")

# === COMPANY CANDIDATES ===
company_candidates = []

def add_company_candidate(val, score, why):
    cleaned = clean_company(val)
    if not cleaned:
        return
    if is_probable_location(cleaned):
        debug(f"Rejected company candidate '{cleaned}' (looks like location) from {why}")
        return
    company_candidates.append((cleaned, score, why))
    debug(f"Company candidate (+{score}): '{cleaned}' via {why}")


# Explicit labels
m = re.search(r"(?:Job\s*Title|Position|Role|Title)\s*[:\-–]\s*([^\n]+)", job_description, flags=re.IGNORECASE)
if m: add_title_candidate(m.group(1), 6, "label")

# Looking for a/an ...
m = re.search(r"(?i)looking\s+for\s+a[n]?\s+([^\n]+)", job_description)
if m: add_title_candidate(m.group(1), 5, "looking for a/an")

# "As a/an ..." pattern ...
m = re.search(r"(?i)\bas\s+a[n]?\s+([^\n,]+)", job_description)
if m:
    candidate = m.group(1).strip()

    # Only keep if it starts with a capital letter and isn't a generic filler phrase
    generic_phrases = {
        "member of the team", "part of the team", "part of our team",
        "member of our company", "part of the company", "team member"
    }
    if candidate and candidate[0].isupper() and candidate.lower() not in generic_phrases:
        add_title_candidate(candidate, 5, "as a/an")


# Early lines: separators and title-cased
lines = job_description.splitlines()

for idx, ln in enumerate(lines[:5]):
    ln = ln.strip()
    if not ln:
        continue

    # Separator split logic stays the same
    for sep in [" - ", " – ", " — ", ":", " | "]:
        if sep in ln:
            left, right = ln.split(sep, 1)
            if left and left[0].isupper():
                add_title_candidate(left, 4, f"before '{sep.strip()}'")
            if right and right[0].isupper():
                add_title_candidate(right, 3, f"after '{sep.strip()}'")

    # Title-cased line detection with lookahead
    if ln and ln[0].isupper() and looks_title_cased(ln) and not is_probable_location(ln):
        company, job_title, needs_lookahead = extract_company_and_title_from_line(ln)

        if needs_lookahead:
            found_company = find_company_in_next_lines(lines, idx)
            if found_company:
                add_company_candidate(found_company, 4, "lookahead after 'Join us'")
        
        if company:
            add_company_candidate(company, 4, "title-cased line (extracted)")

        if job_title:
            cleaned_title = clean_title(job_title)
            if cleaned_title and not is_probable_location(cleaned_title):
                add_title_candidate(cleaned_title, 3, "title-cased line (extracted)")
        else:
            cleaned_line = clean_title(ln)
            if cleaned_line and not is_probable_location(cleaned_line):
                add_title_candidate(cleaned_line, 3, "title-cased line")


# Title @/at Company
m = re.search(r"^([A-Z][A-Za-z0-9/&\-\s]+?)\s*@\s*[A-Z][^\n]+", job_description, flags=re.MULTILINE)
if m: add_title_candidate(m.group(1), 5, "Title @ Company")
m = re.search(r"^([A-Z][A-Za-z0-9/&\-\s]+?)\s+at\s+[A-Z][^\n]+", job_description, flags=re.MULTILINE|re.IGNORECASE)
if m: add_title_candidate(m.group(1), 4, "Title at Company")

def pick_best_title():
    best, best_score = None, -999
    for val, base, why in title_candidates:
        cleaned = clean_title(val)
        if cleaned.strip().lower() in BANNED_TITLES:
            debug(f"Rejecting title '{cleaned}' (banned)")
            continue
        score = base
        if has_role_keyword(val): score += 3
        if looks_title_cased(val): score += 2
        if COMMON_FILLER_START.match(val): score -= 4
        if len(val.split()) > 10: score -= 3
        if not val[0].isupper(): score -= 2
        if any(pref in cleaned.lower() for pref in PREFERRED_TITLES):
            score += 5
            debug(f"Boosting '{cleaned}' for matching preferred role")
        # Penalize if cleaning removed too much (likely junky phrase)
        if len(cleaned) <= max(3, len(val) * 0.5): score -= 1
        debug(f"Title score {score} for '{cleaned}' (from '{val}' via {why})")
        if score > best_score:
            best_score, best = score, cleaned
    return best

best_title = pick_best_title()

# Explicit labels
m = re.search(r"(?:Company|Employer|Organization|Hiring\s*Organization)\s*[:\-–]\s*([^\n]+)", job_description, flags=re.IGNORECASE)
if m: add_company_candidate(m.group(1), 7, "label")

# @ Company
for m in re.finditer(r"@\s*([A-Z][A-Za-z0-9&.,' ]+)", job_description):
    add_company_candidate(m.group(1), 6, "@ Company")

# at Company (avoid common non-company phrases)
for m in re.finditer(r"\bat\s+(?!the\b|scale\b|least\b|most\b)([A-Z][A-Za-z0-9&.,' ]+)", job_description, flags=re.IGNORECASE):
    add_company_candidate(m.group(1), 5, "at Company")

# Careers at / Join
for ln in job_description.splitlines()[:12]:
    ln = ln.strip()
    if not ln:
        continue
    m = re.search(r"(?i)(?:careers\s+at|join)\s+([A-Z][A-Za-z0-9&.,' ]+)", ln)
    if m: add_company_candidate(m.group(1), 4, "careers/join line")

# Proper-noun-like early lines without verbs and without role keywords
for ln in job_description.splitlines()[:8]:
    ln = ln.strip()
    if not ln:
        continue
    if not NON_NAME_VERBS.search(ln) and looks_title_cased(ln) and not has_role_keyword(ln) and not is_probable_location(ln):
        add_company_candidate(ln, 3, "proper-noun line")

def is_probable_company(name):
    if not name or len(name) < 2:
        return False
    if EMAIL_OR_URL.search(name):
        return False
    if COMMON_FILLER_START.match(name):
        return False
    if NON_NAME_VERBS.search(name):
        return False
    if is_probable_location(name):
        return False
    words = name.split()
    if len(words) > 6:
        return False
    if not any(w and w[0].isupper() for w in words):
        return False
    if not re.search(r"[A-Za-z]", name):
        return False
    return True

def company_score(name):
    score = 0
    if re.search(r"\b(Inc|LLC|Ltd|Limited|Corporation|Corp|GmbH|PLC|Pte|BV|S\.A\.|SAS)\b\.?", name):
        score += 3
    if looks_title_cased(name):
        score += 2
    if len(name.split()) > 4:
        score -= 2
    if TRAILING_DEPT.search(name):
        score -= 2
    if name.isupper() or name.islower():
        score -= 1
    return score

def pick_best_company():
    best, best_score = None, -999
    for val, base, why in company_candidates:
        if not is_probable_company(val):
            debug(f"Reject company '{val}' (fails plausibility) from {why}")
            continue
        score = base + company_score(val)
        # Penalize trailing simple location fragment (e.g., ", Austin")
        if re.search(r",\s*[A-Z][a-z]+$", val):
            score -= 2
        # Prefer shorter clean names
        if len(val) > 40:
            score -= 3
        debug(f"Company score {score} for '{val}' via {why}")
        if score > best_score:
            best_score, best = score, val
    return best

best_company = pick_best_company()

# === CONFIRM/EDIT ===
if not best_title:
    best_title = input("Enter the job title: ").strip()
else:
    user = input(f"Detected job title: '{best_title}'. Press Enter to accept or type a correction: ").strip()
    if user:
        best_title = user

if not best_company:
    best_company = input("Enter the company name: ").strip()
else:
    user = input(f"Detected company: '{best_company}'. Press Enter to accept or type a correction: ").strip()
    if user:
        best_company = user

# === CHOOSE OUTPUT FORMAT ===
fmt = input(f"Save format [{DEFAULT_FORMAT}] (txt/docx): ").strip().lower() or DEFAULT_FORMAT
if fmt not in {"txt", "docx"}:
    print("Invalid choice. Falling back to .txt.")
    fmt = "txt"

# === SAVE FILE ===
timestamp = datetime.now().strftime("%Y-%m-%d_%H-%M")
file_name_base = f"{sanitize_filename(best_company)} - {sanitize_filename(best_title)} - {timestamp}"
file_path = os.path.join(SAVE_FOLDER, f"{file_name_base}.{fmt}")

if fmt == "txt":
    with open(file_path, "w", encoding="utf-8") as f:
        f.write(f"{best_title}\n")
        f.write(f"Company: {best_company}\n\n")
        f.write(job_description)
else:
    try:
        from docx import Document  # lazy import
        doc = Document()
        doc.add_heading(best_title, level=1)
        doc.add_paragraph(f"Company: {best_company}")
        doc.add_paragraph("")  # blank line
        doc.add_paragraph(job_description)
        doc.save(file_path)
    except Exception as e:
        print(f"Could not save as .docx ({e}). Saving as .txt instead.")
        file_path = os.path.join(SAVE_FOLDER, f"{file_name_base}.txt")
        with open(file_path, "w", encoding="utf-8") as f:
            f.write(f"{best_title}\n")
            f.write(f"Company: {best_company}\n\n")
            f.write(job_description)

print(f"Job description saved to: {file_path}")